#test-docx.py_1
from docx import Document
import wikipedia

#ใส่ฟังชั่นให้มัน_7
def wiki(keyword,lang = 'th'):
     #3
     wikipedia.set_lang(lang)
     #summary สำหรับบทความสรุป_2
     data = wikipedia.summary(keyword)
     #ข้อมูล page + content_6 บทความทั้งหน้า
     data2 = wikipedia.page(keyword)
     data2 = data2.content
     #สร้างไฟล์ word ใน python_4
     doc = Document()
     doc.add_heading(keyword,0)
     doc.add_paragraph(data2)
     #save_5
     doc.save(keyword +'.docx')
     print('สร้างไฟล์สำเร็จ')


wiki('united state of america','en')

