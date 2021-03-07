#GUIWIKI.py_1
import wikipedia

#python to docx _9
from docx import Document


#python to docx _10
def wiki(keyword,lang = 'th'):
     #3
     wikipedia.set_lang(lang)
     #summary สำหรับบทความสรุป_2
     data = wikipedia.summary(keyword)
     #ข้อมูล page + content_6 บทความทั้งหน้า
     data2 = wikipedia.page(keyword)
     data2 = data2.content
     #สร้างไฟล์ word ใน python_4
     doc = Document()
     doc.add_heading(keyword,0)
     doc.add_paragraph(data2)
     #save_5
     doc.save(keyword +'.docx')
     print('สร้างไฟล์สำเร็จ')



#เปลี่ยนภาษาเป็นภาษาไทย_2
#จากไลบรารีชื่อ tkinter, * คือให้ดึงความสามารถหลักมาทั้งหมด_2.1
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox #10


#ประกาศGUI สร้างหน้าต่างหลัก_3
GUI = Tk() #จะต้องคู่กับGUI.mainloop() ถ้าไม่มีจะรันไม่ได้
GUI.title('โปรแกรม wiki')
GUI.geometry('400x300')

#config ตัวหนังสื_6
FONT1 = ('Angsana New',15)

#คำอธิบาย_7
L = ttk.Label(GUI, text = 'ค้นหาบทความ',font = FONT1)
L.pack()


#สร้างช่องค้นหาข้อมูล Entry (ช่องกรอกข้อความ)_4
#4.2
v_search = StringVar() #กล่องกำหรับเก็บคีย์เวิร์ด
#4.1
E1 = ttk.Entry(GUI,textvariable = v_search,font = FONT1,width=40)
#4.3ระยะห่าง ช่องว่าง 
E1.pack(pady=20)

#ฟังชั่นปุ่มค้นหา_8
def search():
     keyword = v_search.get()  #.get () คือดึงข้อมูลเข้ามา
     try:
          language = v_radio.get() #th/en/zh_16
          wiki(keyword, language)
          messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อยแล้ว')
     except:
          #หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
          messagebox.showwarning('keyword Error', 'กรุณากรอกคำค้นหาใหม่')
          

     
    # print(wikipedia.search(keyword))
 #   result = wikipedia.summary(keyword ,)
   #  print(result)



#ปุ๋มค้นหา_5
#ipadx ขยายปุ่มแนวนอน
#ipady ขยายปุ่มแนวตั้ง 
B1 = ttk.Button(GUI,text = 'search' , command = search)
B1.pack(ipadx=20,ipady=10)

# เลือกภาษา_13
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar() #ช่องเก็บข้อมูล

RB1 = ttk.Radiobutton(F1,text = 'ภาษาไทย',variable = v_radio,value ='th') #13
RB2 = ttk.Radiobutton(F1,text = 'อังกฤษ',variable = v_radio,value ='en') #13
RB3 = ttk.Radiobutton(F1,text = 'จีน',variable = v_radio,value ='zn') #13
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย_14

RB1.grid(row =0,column=0) #ทำให้เป็นแนวนอน_15
RB2.grid(row =0,column=1)
RB3.grid(row =0,column=2)


GUI.mainloop()
